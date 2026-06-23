import os, sqlite3, streamlit as st
from groq import Groq
from dotenv import load_dotenv
import urllib.request, urllib.parse, json, re
from io import BytesIO

DB_PATH     = './ptsp_docs.db'
DB_DRIVE_ID = '11aXRiR5N7XDsZuGSKZmp_eVLmFyN0Z25'

if not os.path.exists(DB_PATH):
    with st.spinner('Mengunduh database (hanya sekali, mohon tunggu)...'):
        import gdown
        url = f'https://drive.google.com/uc?id={DB_DRIVE_ID}'
        gdown.download(url, DB_PATH, quiet=False)
    st.success('Database berhasil diunduh!')

try:
    GROQ_API_KEY = st.secrets['GROQ_API_KEY']
    GROQ_MODEL   = st.secrets.get('GROQ_MODEL', 'llama-3.3-70b-versatile')
except Exception:
    load_dotenv()
    GROQ_API_KEY = os.getenv('GROQ_API_KEY')
    GROQ_MODEL   = os.getenv('GROQ_MODEL', 'llama-3.3-70b-versatile')

groq_client = Groq(api_key=GROQ_API_KEY)

st.set_page_config(page_title='AI Assistant — PTSP', page_icon='🤖', layout='wide')

st.markdown("""
<style>
[data-testid="stSidebar"] { background: #1E3A5F; }
[data-testid="stSidebar"] * { color: white !important; }
[data-testid="stSidebar"] .stButton > button {
    background: rgba(255,255,255,0.1); color: white !important;
    border: 1px solid rgba(255,255,255,0.2); border-radius: 8px;
    text-align: left; font-size: 13px;
}
[data-testid="stSidebar"] .stButton > button:hover { background: rgba(255,255,255,0.2); }
.source-box { background:#f0f7ff; border-left:3px solid #1E3A5F;
    padding:8px 12px; border-radius:4px; font-size:13px; margin-top:8px; }
.web-box { background:#f0fff4; border-left:3px solid #1A5C3A;
    padding:8px 12px; border-radius:4px; font-size:13px; margin-top:4px; }
</style>
""", unsafe_allow_html=True)

KAMUS = {
    'pompa':'pump','pipa':'pipe','katup':'valve','motor':'motor',
    'listrik':'electrical','pelumas':'lubrication','oli':'oil',
    'tekanan':'pressure','suhu':'temperature','aliran':'flow',
    'perawatan':'maintenance','perbaikan':'repair','jarak':'gap',
    'celah':'gap','ukuran':'dimension','prosedur':'procedure',
    'spesifikasi':'specification','kalibrasi':'calibration',
    'instrumen':'instrument','kompresor':'compressor','turbin':'turbine',
    'generator':'generator','kabel':'cable','keselamatan':'safety',
    'gangguan':'fault','getaran':'vibration','inspeksi':'inspection',
    'pengujian':'testing','penggantian':'replacement','tangki':'tank',
    'pendingin':'cooling','pemanas':'heating','plat':'plate',
    'stator':'stator','rotor':'rotor','shaft':'shaft','roller':'roller',
    'separator':'separator','fan':'fan','belt':'belt','gearbox':'gearbox',
    'bearing':'bearing','seal':'seal','coupling':'coupling',
    'impeller':'impeller','nozzle':'nozzle','filter':'filter',
    'mill':'mill','kiln':'kiln','crusher':'crusher','blower':'blower',
    'conveyor':'conveyor','elevator':'elevator','silo':'silo',
    'standar':'standard','formulir':'form','pemeriksaan':'inspection',
    'pengukuran':'measurement','toleransi':'tolerance','raw':'raw',
    'finish':'finish','cement':'cement','clinker':'clinker',
    'qc':'qc','form':'form','checklist':'checklist','drawing':'drawing',
    'gambar':'drawing','part':'part','daftar':'list',
}

EQUIPMENT_CODES = {
    'raw mill':    ['6R1'],
    'coal mill':   ['6K1'],
    'cement mill': ['6Z1'],
    'kiln':        ['6W1'],
    'cooler':      ['6W1', 'crossbar'],
}

def translate_query(query):
    return [KAMUS.get(w.strip('.,?!()'), w.strip('.,?!()'))
            for w in query.lower().split()]

def detect_equipment_code(query):
    query_lower = query.lower()
    for name, codes in EQUIPMENT_CODES.items():
        if name in query_lower:
            return codes
    return []

def web_search(query, max_results=3):
    try:
        encoded = urllib.parse.quote(f"{query} cement plant technical specification")
        url     = f"https://api.duckduckgo.com/?q={encoded}&format=json&no_html=1&skip_disambig=1"
        req     = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
        with urllib.request.urlopen(req, timeout=5) as r:
            data = json.loads(r.read().decode())
        results = []
        if data.get('AbstractText'):
            results.append({'title': data.get('Heading','Ref'),
                           'snippet': data['AbstractText'][:400],
                           'url': data.get('AbstractURL','')})
        for topic in data.get('RelatedTopics', [])[:max_results]:
            if isinstance(topic, dict) and topic.get('Text'):
                results.append({'title': topic['Text'][:60],
                               'snippet': topic['Text'][:300],
                               'url': topic.get('FirstURL','')})
        return results[:max_results]
    except:
        return []

def check_db_has_folder_path():
    try:
        conn = sqlite3.connect(DB_PATH)
        conn.execute('SELECT folder_path FROM docs LIMIT 1')
        conn.close()
        return True
    except:
        return False

def search_db(query, top_k=6):
    try:
        conn   = sqlite3.connect(DB_PATH)
        words  = translate_query(query)
        has_fp = check_db_has_folder_path()
        equipment_codes = detect_equipment_code(query)
        results, seen = [], set()

        if has_fp:
            SELECT = 'SELECT file_name, file_path, content, folder_path FROM docs'
        else:
            SELECT = 'SELECT file_name, file_path, content, "" as folder_path FROM docs'

        def add(rows, boost_codes=None):
            for r in rows:
                fname, fpath, content = r[0], r[1], r[2]
                fpath_folder = r[3] if len(r) > 3 else ''
                if boost_codes:
                    text_check = (fname + ' ' + fpath_folder + ' ' + content[:200]).upper()
                    other_codes = [c for c in ['6R1','6K1','6Z1','6W1'] if c not in boost_codes]
                    has_target = any(c.upper() in text_check for c in boost_codes)
                    has_other  = any(c.upper() in text_check for c in other_codes)
                    if has_other and not has_target:
                        continue
                key = fname + content[:60]
                if key not in seen:
                    seen.add(key)
                    results.append({
                        'file_name':   fname,
                        'file_path':   fpath,
                        'content':     content,
                        'folder_path': fpath_folder
                    })

        if has_fp and equipment_codes:
            for code in equipment_codes:
                try:
                    rows = conn.execute(
                        f"{SELECT} WHERE folder_path LIKE ? OR file_name LIKE ? OR content LIKE ? LIMIT ?",
                        (f'%{code}%', f'%{code}%', f'%{code}%', top_k * 2)
                    ).fetchall()
                    add(rows, boost_codes=equipment_codes)
                except: pass

        folder_keywords = ['qc', 'qc-form', 'qcform', 'drawing', 'manual',
                          'spesifikasi', 'specification', 'part list', 'partlist']
        query_lower = query.lower()
        folder_hint = next((k for k in folder_keywords if k in query_lower), None)

        if has_fp and folder_hint:
            folder_map = {
                'qc': 'QC-FORM', 'qc-form': 'QC-FORM', 'qcform': 'QC-FORM',
                'drawing': 'Drawing', 'manual': 'Manual & Spesifikasi',
                'spesifikasi': 'Manual & Spesifikasi',
                'specification': 'Manual & Spesifikasi',
                'part list': 'Part List', 'partlist': 'Part List',
            }
            target_folder = folder_map.get(folder_hint, '')
            if target_folder:
                try:
                    rows = conn.execute(
                        f"{SELECT} WHERE folder_path LIKE ? ORDER BY rank LIMIT ?",
                        (f'%{target_folder}%', top_k * 2)
                    ).fetchall()
                    add(rows, boost_codes=equipment_codes if equipment_codes else None)
                except: pass

        for i in range(len(words) - 2):
            phrase = ' '.join(words[i:i+3])
            try:
                rows = conn.execute(
                    f"{SELECT} WHERE docs MATCH ? ORDER BY rank LIMIT ?",
                    (f'"{phrase}"', top_k)
                ).fetchall()
                add(rows, boost_codes=equipment_codes if equipment_codes else None)
            except: pass

        for i in range(len(words) - 1):
            phrase = ' '.join(words[i:i+2])
            try:
                rows = conn.execute(
                    f"{SELECT} WHERE docs MATCH ? ORDER BY rank LIMIT ?",
                    (f'"{phrase}"', top_k)
                ).fetchall()
                add(rows, boost_codes=equipment_codes if equipment_codes else None)
            except: pass

        if len(results) < 3:
            for word in sorted(set(words), key=len, reverse=True)[:4]:
                if len(word) < 3: continue
                try:
                    rows = conn.execute(
                        f"{SELECT} WHERE docs MATCH ? ORDER BY rank LIMIT ?",
                        (word, top_k)
                    ).fetchall()
                    add(rows, boost_codes=equipment_codes if equipment_codes else None)
                except: pass

        conn.close()
        return results[:top_k]
    except Exception as e:
        return []

@st.cache_resource
def get_db_stats():
    if not os.path.exists(DB_PATH): return 0, 0
    conn   = sqlite3.connect(DB_PATH)
    files  = conn.execute('SELECT COUNT(*) FROM indexed_files').fetchone()[0]
    chunks = conn.execute('SELECT COUNT(*) FROM docs').fetchone()[0]
    conn.close()
    return files, chunks

def generate_excel_answer(answer_text, sources, query):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment

    wb = Workbook()
    ws = wb.active
    ws.title = "Hasil AI Assistant"

    ws['A1'] = 'Pertanyaan'
    ws['A2'] = query
    ws['A4'] = 'Jawaban'
    ws['A5'] = answer_text

    for cell in ['A1', 'A4']:
        ws[cell].font = Font(bold=True, size=12, color='FFFFFF')
        ws[cell].fill = PatternFill(start_color='1E3A5F', fill_type='solid')

    ws.column_dimensions['A'].width = 100
    ws['A2'].alignment = Alignment(wrap_text=True)
    ws['A5'].alignment = Alignment(wrap_text=True)
    ws.row_dimensions[5].height = 200

    row = 7
    ws[f'A{row}'] = 'Sumber Dokumen'
    ws[f'A{row}'].font = Font(bold=True, size=12, color='FFFFFF')
    ws[f'A{row}'].fill = PatternFill(start_color='1A5C3A', fill_type='solid')
    row += 1
    for s in sources:
        ws[f'A{row}'] = f"• {s['file_name']} ({s.get('folder_path','')})"
        row += 1

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf

def generate_docx(answer_text, sources, query):
    from docx import Document as DocxDoc

    doc = DocxDoc()
    doc.add_heading('AI Assistant — Unit Evaluasi Proses dan Energi PTSP', level=1)

    doc.add_heading('Pertanyaan', level=2)
    doc.add_paragraph(query)

    doc.add_heading('Jawaban', level=2)
    doc.add_paragraph(answer_text)

    doc.add_heading('Sumber Dokumen', level=2)
    for s in sources:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{s['file_name']} ({s.get('folder_path','')})")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def classify_intent(query, equipment_codes):
    prompt = f"""Analisis pertanyaan berikut dan tentukan klasifikasinya dalam format JSON murni (tanpa markdown):

Pertanyaan: "{query}"

Tentukan:
1. "mode": "rekap_dokumen" jika user minta KUMPULAN/SEMUA/KESELURUHAN/REKAP dokumen dari satu kategori (misal: "semua QC Form raw mill", "rekap drawing kiln", "kumpulan part list separator"). Pilih "jawab_singkat" jika user hanya tanya 1 hal spesifik atau minta penjelasan/analisis.
2. "format_file": "excel" jika user minta/cocok untuk Excel, "word" jika cocok dokumen Word, "tidak_ada" jika tidak perlu file.
3. "target_folder": pilih SALAH SATU dari ["QC-FORM", "Drawing", "Manual & Spesifikasi", "Part List", null] sesuai konteks pertanyaan. null jika tidak spesifik ke satu folder.

Jawab HANYA dengan JSON, contoh:
{{"mode": "rekap_dokumen", "format_file": "excel", "target_folder": "QC-FORM"}}"""

    try:
        completion = groq_client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[{'role': 'user', 'content': prompt}],
            max_tokens=150,
            temperature=0
        )
        raw = completion.choices[0].message.content.strip()
        raw = re.sub(r'^```json\s*|\s*```$', '', raw, flags=re.MULTILINE).strip()
        result = json.loads(raw)
        return {
            'mode': result.get('mode', 'jawab_singkat'),
            'format_file': result.get('format_file', 'tidak_ada'),
            'target_folder': result.get('target_folder')
        }
    except Exception:
        return {'mode': 'jawab_singkat', 'format_file': 'tidak_ada', 'target_folder': None}

def list_files_for_folder(target_folder, equipment_codes, limit=300):
    conn = sqlite3.connect(DB_PATH)
    results, seen_files = [], set()
    try:
        if target_folder:
            if equipment_codes:
                for code in equipment_codes:
                    rows = conn.execute('''
                        SELECT DISTINCT file_name, folder_path, file_path
                        FROM docs_content
                        WHERE folder_path LIKE ?
                        AND (file_name LIKE ? OR folder_path LIKE ? OR content LIKE ?)
                        LIMIT ?
                    ''', (f'%{target_folder}%', f'%{code}%', f'%{code}%', f'%{code}%', limit)).fetchall()
                    for fname, fpath, link in rows:
                        if fname not in seen_files:
                            seen_files.add(fname)
                            results.append({'file_name': fname, 'folder_path': fpath, 'file_path': link or ''})
            else:
                rows = conn.execute('''
                    SELECT DISTINCT file_name, folder_path, file_path
                    FROM docs_content
                    WHERE folder_path LIKE ?
                    LIMIT ?
                ''', (f'%{target_folder}%', limit)).fetchall()
                for fname, fpath, link in rows:
                    if fname not in seen_files:
                        seen_files.add(fname)
                        results.append({'file_name': fname, 'folder_path': fpath, 'file_path': link or ''})
        elif equipment_codes:
            for code in equipment_codes:
                rows = conn.execute('''
                    SELECT DISTINCT file_name, folder_path, file_path
                    FROM docs_content
                    WHERE file_name LIKE ? OR folder_path LIKE ? OR content LIKE ?
                    LIMIT ?
                ''', (f'%{code}%', f'%{code}%', f'%{code}%', limit)).fetchall()
                for fname, fpath, link in rows:
                    if fname not in seen_files:
                        seen_files.add(fname)
                        results.append({'file_name': fname, 'folder_path': fpath, 'file_path': link or ''})
    except Exception:
        pass
    conn.close()
    return results

def generate_excel_recap(files, query, folder_label=""):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment

    conn = sqlite3.connect(DB_PATH)
    wb = Workbook()
    ws = wb.active
    ws.title = "Rekap Dokumen"

    ws['A1'] = f'Rekap Dokumen — {folder_label or query}'
    ws.merge_cells('A1:E1')
    ws['A1'].font = Font(bold=True, size=14, color='FFFFFF')
    ws['A1'].fill = PatternFill(start_color='1E3A5F', fill_type='solid')
    ws['A1'].alignment = Alignment(horizontal='center')

    headers = ['No', 'Nama File', 'Folder', 'Isi / Konten', 'Link Dokumen']
    for i, h in enumerate(headers, start=1):
        cell = ws.cell(row=3, column=i, value=h)
        cell.font = Font(bold=True, color='FFFFFF')
        cell.fill = PatternFill(start_color='1A5C3A', fill_type='solid')

    current_row = 4
    for idx, f in enumerate(files, start=1):
        chunks = conn.execute('''
            SELECT content FROM docs_content
            WHERE file_name = ?
            ORDER BY CAST(chunk_id AS INTEGER)
        ''', (f['file_name'],)).fetchall()
        full_content = ' '.join(c[0] for c in chunks).strip()
        full_content = re.sub(r'\s+', ' ', full_content)[:2000]

        ws.cell(row=current_row, column=1, value=idx)
        ws.cell(row=current_row, column=2, value=f['file_name'])
        ws.cell(row=current_row, column=3, value=f.get('folder_path',''))
        c = ws.cell(row=current_row, column=4, value=full_content)
        c.alignment = Alignment(wrap_text=True, vertical='top')
        ws.cell(row=current_row, column=5, value=f.get('file_path',''))
        current_row += 1

    conn.close()
    ws.column_dimensions['A'].width = 6
    ws.column_dimensions['B'].width = 45
    ws.column_dimensions['C'].width = 20
    ws.column_dimensions['D'].width = 80
    ws.column_dimensions['E'].width = 50

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf

def generate_docx_recap(files, query, folder_label=""):
    from docx import Document as DocxDoc

    conn = sqlite3.connect(DB_PATH)
    doc = DocxDoc()
    doc.add_heading(f'Rekap Dokumen — {folder_label or query}', level=1)

    for idx, f in enumerate(files, start=1):
        chunks = conn.execute('''
            SELECT content FROM docs_content
            WHERE file_name = ?
            ORDER BY CAST(chunk_id AS INTEGER)
        ''', (f['file_name'],)).fetchall()
        full_content = ' '.join(c[0] for c in chunks).strip()
        full_content = re.sub(r'\s+', ' ', full_content)[:3000]

        doc.add_heading(f"{idx}. {f['file_name']}", level=2)
        doc.add_paragraph(f"Folder: {f.get('folder_path','')}")
        doc.add_paragraph(full_content)
        if f.get('file_path'):
            doc.add_paragraph(f"Link: {f['file_path']}")
        doc.add_paragraph('')

    conn.close()
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

files_count, chunks_count = get_db_stats()

with st.sidebar:
    st.markdown("### 🤖 AI Assistant PTSP")
    st.markdown("*Unit Evaluasi Proses dan Energi*")
    st.divider()
    st.markdown("**📊 Database**")
    st.metric("File Terindeks", f"{files_count:,}")
    st.metric("Total Chunks", f"{chunks_count:,}")
    st.divider()
    st.markdown("**💡 Contoh Pertanyaan**")
    examples = [
        "QC Form raw mill standar apa saja?",
        "Berikan semua QC Form raw mill dalam excel",
        "Drawing separator raw mill",
        "Rekap part list gearbox raw mill",
        "Hydraulic pressure setting vertical mill",
        "Troubleshooting vibration high pada mill",
        "Manual & spesifikasi OK mill",
    ]
    for ex in examples:
        if st.button(ex, use_container_width=True, key=ex):
            st.session_state['pending'] = ex
    st.divider()
    use_web = st.toggle("🌐 Web search fallback", value=True)
    if st.button("🗑️ Reset chat", use_container_width=True):
        st.session_state.messages = []
        st.rerun()

st.title("🤖 AI Assistant — PTSP")
st.caption("Dokumen: Drawing | Manual & Spesifikasi | Part List | QC-FORM")

if 'messages' not in st.session_state:
    st.session_state.messages = []

for msg in st.session_state.messages:
    with st.chat_message(msg['role']):
        st.markdown(msg['content'], unsafe_allow_html=True)

query = st.session_state.pop('pending', None) or st.chat_input(
    "Tanya dalam Bahasa Indonesia atau Inggris..."
)

if query:
    st.session_state.messages.append({'role': 'user', 'content': query})
    with st.chat_message('user'):
        st.markdown(query)

    with st.chat_message('assistant'):
        with st.spinner('Memahami konteks pertanyaan...'):
            equipment_codes = detect_equipment_code(query)
            intent = classify_intent(query, equipment_codes)

        if intent['mode'] == 'rekap_dokumen':
            with st.spinner('Mengambil dan menyusun dokumen...'):
                files = list_files_for_folder(intent['target_folder'], equipment_codes)

            if files:
                folder_label = intent['target_folder'] or (", ".join(equipment_codes) if equipment_codes else query)
                answer = f"Ditemukan **{len(files)} dokumen** sesuai permintaan Anda di folder **{intent['target_folder'] or 'seluruh database'}**. Isi lengkap setiap dokumen sudah disusun ke dalam file di bawah ini."
                st.markdown(answer)

                with st.spinner(f'Mengekstrak isi {len(files)} dokumen...'):
                    if intent['format_file'] == 'word':
                        file_buf = generate_docx_recap(files, query, folder_label)
                        st.download_button(
                            label=f"📄 Download Word — {len(files)} Dokumen",
                            data=file_buf,
                            file_name=f"rekap_{folder_label.replace(' ','_')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"docx_recap_{len(st.session_state.messages)}"
                        )
                    else:
                        file_buf = generate_excel_recap(files, query, folder_label)
                        st.download_button(
                            label=f"📊 Download Excel — {len(files)} Dokumen",
                            data=file_buf,
                            file_name=f"rekap_{folder_label.replace(' ','_')}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            key=f"excel_recap_{len(st.session_state.messages)}"
                        )

                html = f'<div class="source-box">📄 <b>Folder:</b> {intent["target_folder"] or "Seluruh database"}<br><b>Total dokumen:</b> {len(files)}</div>'
                st.markdown(html, unsafe_allow_html=True)
            else:
                answer = "Mohon maaf, tidak ditemukan dokumen yang sesuai dengan permintaan Anda."
                st.markdown(answer)

            st.session_state.messages.append({'role': 'assistant', 'content': answer})

        else:
            with st.spinner('Mencari...'):
                docs        = search_db(query)
                web_results = []

                if use_web and len(docs) < 2:
                    web_query   = ' '.join(translate_query(query))
                    web_results = web_search(web_query)

                ctx_parts = []
                if docs:
                    ctx_parts.append("=== DOKUMEN INTERNAL PTSP ===")
                    for d in docs:
                        header = f"[{d['file_name']}]"
                        if d.get('folder_path'):
                            header += f" (📁 {d['folder_path']})"
                        ctx_parts.append(f"{header}\n{d['content']}")
                else:
                    ctx_parts.append("=== DOKUMEN INTERNAL ===\nTidak ada dokumen relevan.")

                if web_results:
                    ctx_parts.append("\n=== REFERENSI TEKNIS INTERNET ===")
                    for w in web_results:
                        ctx_parts.append(f"[{w['title']}]\n{w['snippet']}")

                context = '\n\n---\n\n'.join(ctx_parts)

                system_prompt = """Anda adalah AI Assistant teknis untuk Unit Evaluasi Proses dan Energi PTSP (pabrik semen Indarung VI).

KODE EQUIPMENT (PENTING — jangan tertukar):
- 6R1 = Raw Mill (OK Mill) — vertical roller mill untuk material mentah
- 6K1 = Coal Mill — vertical mill untuk batu bara
- 6Z1 = Cement Mill — vertical mill untuk semen
- 6W1 = Kiln & Cooler — rotary kiln dan crossbar cooler

Equipment dengan kode berbeda (6R1 vs 6K1 vs 6Z1 vs 6W1) adalah MESIN YANG BERBEDA
meskipun sama-sama 'Vertical Mill'. JANGAN campurkan informasi dari kode equipment
yang berbeda.

STRUKTUR FOLDER DATABASE:
- Drawing        → gambar teknis, P&ID, GA drawing
- Manual & Spesifikasi → manual operasi, spesifikasi teknik
- Part List      → daftar part/spare part equipment
- QC-FORM        → form pemeriksaan & quality control tiap equipment

ATURAN:
1. Prioritaskan DOKUMEN INTERNAL. Sebutkan nama file dan folder sumbernya.
2. Referensi internet hanya untuk konteks teknis jika dokumen tidak ada.
3. Jawab Bahasa Indonesia, teknis dan formal.
4. Terjemahkan dokumen Inggris ke Indonesia dalam jawaban.
5. Jika dokumen tidak ditemukan, katakan terus terang."""

                try:
                    completion = groq_client.chat.completions.create(
                        model=GROQ_MODEL,
                        messages=[
                            {'role': 'system', 'content': system_prompt},
                            {'role': 'user', 'content': f"Konteks:\n{context}\n\nPertanyaan: {query}"}
                        ],
                        max_tokens=2000,
                        temperature=0.1
                    )
                    answer = completion.choices[0].message.content
                    st.markdown(answer)

                    if intent['format_file'] in ('excel', 'word') and docs:
                        col1, col2 = st.columns(2)
                        with col1:
                            if intent['format_file'] == 'excel':
                                buf = generate_excel_answer(answer, docs, query)
                                st.download_button("📊 Download Excel", data=buf,
                                    file_name=f"jawaban_{query[:30].replace(' ','_')}.xlsx",
                                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                    key=f"excel_{len(st.session_state.messages)}")
                        with col2:
                            if intent['format_file'] == 'word':
                                buf = generate_docx(answer, docs, query)
                                st.download_button("📄 Download Word", data=buf,
                                    file_name=f"jawaban_{query[:30].replace(' ','_')}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key=f"docx_{len(st.session_state.messages)}")

                    if docs:
                        html = '<div class="source-box">📄 <b>Sumber:</b><br>'
                        for d in docs:
                            label = d['file_name']
                            if d.get('folder_path'):
                                label += f" <i>({d['folder_path']})</i>"
                            if d['file_path']:
                                html += f'• <a href="{d["file_path"]}" target="_blank">{label}</a><br>'
                            else:
                                html += f'• {label}<br>'
                        html += '</div>'
                        st.markdown(html, unsafe_allow_html=True)

                    if web_results:
                        whtml = '<div class="web-box">🌐 <b>Referensi Internet:</b><br>'
                        for w in web_results:
                            if w['url']:
                                whtml += f'• <a href="{w["url"]}" target="_blank">{w["title"][:70]}</a><br>'
                        whtml += '</div>'
                        st.markdown(whtml, unsafe_allow_html=True)

                    full = answer
                    if docs:
                        full += '\n\n**Sumber:** ' + ', '.join(
                            f"{d['file_name']} ({d.get('folder_path','')})" for d in docs
                        )
                    st.session_state.messages.append({'role': 'assistant', 'content': full})

                except Exception as e:
                    err = f'❌ Error: {e}'
                    st.error(err)
                    st.session_state.messages.append({'role': 'assistant', 'content': err})